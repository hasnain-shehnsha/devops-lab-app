from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── helpers ──────────────────────────────────────────────────────────────────

def set_margins(doc, top=1, bottom=1, left=1.2, right=1.2):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

def heading(text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
        run.font.size = Pt(14 if level == 1 else 12)
    return p

def body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'E8E8E8')
    p._p.pPr.append(shading)
    return p

def screenshot_box(label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ SCREENSHOT: {label} ]')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'FFF3CD')
    p._p.pPr.append(shading)
    # dashed border
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'dashed')
        bdr.set(qn('w:sz'), '6')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), 'FF0000')
        pBdr.append(bdr)
    p._p.pPr.append(pBdr)
    doc.add_paragraph()   # spacer

def divider():
    doc.add_paragraph('─' * 80)

set_margins(doc)

# ── COVER PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Lab Exam Report')
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('DevOps & Cloud — 3 Hours Exam')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ('Student Name', 'Husnain Shehnsha'),
    ('Email',        'fakhirhassanllc@gmail.com'),
    ('Total Marks',  '50'),
    ('Date',         '21 April 2026'),
]
for i, (k, v) in enumerate(data):
    info_table.cell(i, 0).text = k
    info_table.cell(i, 1).text = v
    for cell in info_table.rows[i].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION B — App Development  (do before A so files exist for A)
# ═══════════════════════════════════════════════════════════════════════════════
heading('Section B: Application Development (6 Marks)')
body('Created a simple Flask web application that returns "Hello DevOps World".', size=11)
doc.add_paragraph()

heading('app.py', level=2)
code_block(
    'from flask import Flask\n\n'
    'app = Flask(__name__)\n\n'
    '@app.route("/")\n'
    'def hello():\n'
    '    return "Hello DevOps World"\n\n'
    'if __name__ == "__main__":\n'
    '    app.run(host="0.0.0.0", port=5000)'
)

heading('requirements.txt', level=2)
code_block('flask')

screenshot_box('Figure B-1: app.py file open in editor showing Flask code')
screenshot_box('Figure B-2: requirements.txt file showing "flask"')
divider()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION A — GitHub
# ═══════════════════════════════════════════════════════════════════════════════
heading('Section A: GitHub / GitLab (8 Marks)')

heading('1. Create Repository', level=2)
body('Created a new public repository named devops-lab-app on GitHub.')
screenshot_box('Figure A-1: GitHub — New repository page with name "devops-lab-app"')
screenshot_box('Figure A-2: GitHub — Repository created successfully (empty repo page)')

heading('2. Add Files & Push Code', level=2)
code_block(
    'cd "/home/husnain-shehnsha/python/dev ops lab mids"\n'
    'git init\n'
    'git add app.py requirements.txt Dockerfile deployment.yaml Jenkinsfile\n'
    'git commit -m "Initial commit: Flask app with Docker and K8s config"\n'
    'git remote add origin https://github.com/YOUR_USERNAME/devops-lab-app.git\n'
    'git branch -m master main\n'
    'git push -u origin main'
)
screenshot_box('Figure A-3: Terminal — git init and git add commands output')
screenshot_box('Figure A-4: Terminal — git commit output showing files committed')
screenshot_box('Figure A-5: Terminal — git push output showing "main" branch pushed')
screenshot_box('Figure A-6: GitHub — Repository page showing app.py and other files')

heading('4. Create Branch feature-update', level=2)
code_block(
    'git checkout -b feature-update\n'
    'git push origin feature-update'
)
screenshot_box('Figure A-7: Terminal — git checkout -b feature-update output')
screenshot_box('Figure A-8: Terminal — git push origin feature-update output')
screenshot_box('Figure A-9: GitHub — Branches page showing both "main" and "feature-update"')

heading('5. Create Pull Request', level=2)
body('Navigated to GitHub → Pull Requests → New Pull Request → base: main ← compare: feature-update → Created PR.')
screenshot_box('Figure A-10: GitHub — New Pull Request page (base: main ← feature-update)')
screenshot_box('Figure A-11: GitHub — Pull Request successfully created')
divider()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION C — Docker
# ═══════════════════════════════════════════════════════════════════════════════
heading('Section C: Docker (10 Marks)')

heading('1. Dockerfile', level=2)
code_block(
    'FROM python:3.9-slim\n\n'
    'WORKDIR /app\n\n'
    'COPY requirements.txt .\n'
    'RUN pip install -r requirements.txt\n\n'
    'COPY app.py .\n\n'
    'EXPOSE 5000\n\n'
    'CMD ["python", "app.py"]'
)
screenshot_box('Figure C-1: Dockerfile open in editor showing all instructions')

heading('2. Build Docker Image', level=2)
code_block('sudo docker build -t devops-lab-app:latest .')
screenshot_box('Figure C-2: Terminal — docker build command running (showing steps 1-6)')
screenshot_box('Figure C-3: Terminal — docker build SUCCESS output "Successfully built <image_id>"')

heading('3. Run Docker Container', level=2)
code_block(
    'sudo docker run -d -p 5000:5000 --name devops-lab-app devops-lab-app:latest\n'
    'sudo docker ps\n'
    'curl http://localhost:5000'
)
screenshot_box('Figure C-4: Terminal — docker run command output showing container ID')
screenshot_box('Figure C-5: Terminal — docker ps showing container running on port 5000')
screenshot_box('Figure C-6: Terminal — curl http://localhost:5000 returning "Hello DevOps World"')
screenshot_box('Figure C-7: Browser — http://localhost:5000 showing "Hello DevOps World"')

heading('4. Push to Docker Hub (Optional)', level=2)
code_block(
    'sudo docker login\n'
    'sudo docker tag devops-lab-app:latest YOUR_USERNAME/devops-lab-app:latest\n'
    'sudo docker push YOUR_USERNAME/devops-lab-app:latest'
)
screenshot_box('Figure C-8 (Optional): Docker Hub — repository page showing pushed image')
divider()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION D — AWS
# ═══════════════════════════════════════════════════════════════════════════════
heading('Section D: AWS Deployment (10 Marks)')

heading('1. Launch EC2 Instance', level=2)
body('Launched Ubuntu 22.04 t2.micro instance on AWS EC2.')
screenshot_box('Figure D-1: AWS Console — EC2 Launch Instance page (Ubuntu 22.04 selected)')
screenshot_box('Figure D-2: AWS Console — EC2 instance running state with Public IP visible')

heading('2. Configure Security Group', level=2)
body('Added inbound rules for ports 22 (SSH), 5000 (Flask app), and 80 (HTTP).')
screenshot_box('Figure D-3: AWS Console — Security Group Inbound Rules showing port 22, 80, 5000 open')

heading('3. Install Docker on EC2', level=2)
code_block(
    'ssh -i devops-mid-lab-key.pem ubuntu@YOUR_EC2_IP\n'
    'sudo apt update && sudo apt install -y docker.io\n'
    'sudo systemctl start docker\n'
    'docker --version'
)
screenshot_box('Figure D-4: Terminal — SSH connection to EC2 successful (ubuntu@ip prompt visible)')
screenshot_box('Figure D-5: Terminal — docker --version output on EC2')

heading('4. Copy Files & Deploy Container', level=2)
code_block(
    '# On LOCAL terminal:\n'
    'scp -i devops-mid-lab-key.pem -r "/home/husnain-shehnsha/python/dev ops lab mids" ubuntu@YOUR_EC2_IP:~/devops-lab-app\n\n'
    '# On EC2 SSH:\n'
    'cd ~/devops-lab-app\n'
    'sudo docker build -t devops-lab-app:latest .\n'
    'sudo docker run -d -p 5000:5000 devops-lab-app:latest\n'
    'sudo docker ps'
)
screenshot_box('Figure D-6: Terminal — scp copy command successful output')
screenshot_box('Figure D-7: Terminal (EC2) — docker build SUCCESS on EC2')
screenshot_box('Figure D-8: Terminal (EC2) — docker ps showing container running')

heading('5. Access via Public IP', level=2)
code_block('curl http://YOUR_EC2_IP:5000')
screenshot_box('Figure D-9: Terminal — curl http://EC2_IP:5000 returning "Hello DevOps World"')
screenshot_box('Figure D-10: Browser — http://EC2_IP:5000 showing "Hello DevOps World"')
divider()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION E — Kubernetes
# ═══════════════════════════════════════════════════════════════════════════════
heading('Section E: Kubernetes (8 Marks)')

heading('deployment.yaml', level=2)
code_block(
    'apiVersion: apps/v1\n'
    'kind: Deployment\n'
    'metadata:\n'
    '  name: devops-lab-app\n'
    'spec:\n'
    '  replicas: 2\n'
    '  selector:\n'
    '    matchLabels:\n'
    '      app: devops-lab-app\n'
    '  template:\n'
    '    metadata:\n'
    '      labels:\n'
    '        app: devops-lab-app\n'
    '    spec:\n'
    '      containers:\n'
    '      - name: devops-lab-app\n'
    '        image: devops-lab-app:latest\n'
    '        imagePullPolicy: Never\n'
    '        ports:\n'
    '        - containerPort: 5000\n'
    '---\n'
    'apiVersion: v1\n'
    'kind: Service\n'
    'metadata:\n'
    '  name: devops-lab-service\n'
    'spec:\n'
    '  selector:\n'
    '    app: devops-lab-app\n'
    '  ports:\n'
    '  - protocol: TCP\n'
    '    port: 80\n'
    '    targetPort: 5000\n'
    '  type: NodePort'
)

heading('1. Start Minikube & Deploy', level=2)
code_block(
    'minikube start --driver=docker\n'
    'minikube image load devops-lab-app:latest\n'
    'kubectl apply -f deployment.yaml'
)
screenshot_box('Figure E-1: Terminal — minikube start output (Done! kubectl is configured)')
screenshot_box('Figure E-2: Terminal — kubectl apply -f deployment.yaml (deployment.apps/devops-lab-app created)')

heading('2. Verify Pods (replicas=2)', level=2)
code_block(
    'kubectl get pods\n'
    'kubectl get deployments\n'
    'kubectl get services'
)
screenshot_box('Figure E-3: Terminal — kubectl get pods showing 2 pods Running')
screenshot_box('Figure E-4: Terminal — kubectl get deployments showing READY 2/2')
screenshot_box('Figure E-5: Terminal — kubectl get services showing devops-lab-service')

heading('3. Expose & Access Service', level=2)
code_block('minikube service devops-lab-service')
screenshot_box('Figure E-6: Terminal — minikube service output with URL')
screenshot_box('Figure E-7: Browser — Kubernetes service URL showing "Hello DevOps World"')
divider()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION F — Jenkins
# ═══════════════════════════════════════════════════════════════════════════════
heading('Section F: Jenkins CI/CD (8 Marks)')

heading('1. Start Jenkins', level=2)
code_block(
    'sudo docker run -d \\\n'
    '  --name jenkins \\\n'
    '  -p 8080:8080 \\\n'
    '  -v jenkins_home:/var/jenkins_home \\\n'
    '  -v /var/run/docker.sock:/var/run/docker.sock \\\n'
    '  jenkins/jenkins:lts\n\n'
    '# Get admin password:\n'
    'sudo docker exec jenkins cat /var/jenkins_home/secrets/initialAdminPassword'
)
screenshot_box('Figure F-1: Terminal — Jenkins Docker container running (docker ps output)')
screenshot_box('Figure F-2: Browser — http://localhost:8080 Jenkins unlock page')
screenshot_box('Figure F-3: Browser — Jenkins dashboard after login')

heading('2. Jenkinsfile', level=2)
code_block(
    "pipeline {\n"
    "    agent any\n"
    "    stages {\n"
    "        stage('Clone Repository') {\n"
    "            steps {\n"
    "                git branch: 'main', url: 'https://github.com/YOUR_USERNAME/devops-lab-app.git'\n"
    "            }\n"
    "        }\n"
    "        stage('Build Docker Image') {\n"
    "            steps {\n"
    "                sh 'docker build -t devops-lab-app:latest .'\n"
    "            }\n"
    "        }\n"
    "        stage('Run Docker Container') {\n"
    "            steps {\n"
    "                sh 'docker stop devops-lab-app || true'\n"
    "                sh 'docker rm devops-lab-app || true'\n"
    "                sh 'docker run -d --name devops-lab-app -p 5000:5000 devops-lab-app:latest'\n"
    "            }\n"
    "        }\n"
    "        stage('Verify') {\n"
    "            steps {\n"
    "                sh 'sleep 3 && curl http://localhost:5000'\n"
    "            }\n"
    "        }\n"
    "    }\n"
    "}"
)

heading('3. Create Pipeline & Connect GitHub', level=2)
body('Steps performed in Jenkins UI:')
body('  1. New Item → Pipeline → Name: devops-lab-pipeline', bold=False)
body('  2. Pipeline → Definition: Pipeline script from SCM', bold=False)
body('  3. SCM: Git → Repository URL: https://github.com/YOUR_USERNAME/devops-lab-app.git', bold=False)
body('  4. Branch: */main → Script Path: Jenkinsfile → Save', bold=False)
body('  5. Build Now', bold=False)
screenshot_box('Figure F-4: Jenkins — New Pipeline item creation page')
screenshot_box('Figure F-5: Jenkins — Pipeline config showing GitHub URL connected')
screenshot_box('Figure F-6: Jenkins — Build Now triggered, pipeline stages visible')
screenshot_box('Figure F-7: Jenkins — All 4 stages GREEN (Clone, Build, Run, Verify)')
screenshot_box('Figure F-8: Jenkins — Console output showing "Hello DevOps World" at end')
divider()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SUMMARY TABLE
# ═══════════════════════════════════════════════════════════════════════════════
heading('Summary — Marks Breakdown')
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
headers = ['Section', 'Task', 'Marks']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(11)

rows_data = [
    ('A', 'GitHub — Repo, Push, Branch, PR',        '8'),
    ('B', 'Flask App — app.py + requirements.txt',   '6'),
    ('C', 'Docker — Dockerfile, Build, Run',         '10'),
    ('D', 'AWS EC2 — Deploy & Access',               '10'),
    ('E', 'Kubernetes — Deploy, Replicas, Service',  '8'),
    ('F', 'Jenkins — Pipeline, GitHub, Docker',      '8'),
    ('',  'Total',                                   '50'),
]
for i, (s, t, m) in enumerate(rows_data, start=1):
    table.cell(i, 0).text = s
    table.cell(i, 1).text = t
    table.cell(i, 2).text = m

doc.add_paragraph()
body('All sections completed successfully. Screenshots attached above each figure label.', bold=True)

# ── Save ──────────────────────────────────────────────────────────────────────
path = '/home/husnain-shehnsha/python/dev ops lab mids/DevOps_Lab_Report.docx'
doc.save(path)
print(f'Document saved: {path}')
